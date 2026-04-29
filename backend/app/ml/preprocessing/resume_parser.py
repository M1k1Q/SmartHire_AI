"""
ml/preprocessing/resume_parser.py
Parses PDF and DOCX files into plain text for NLP processing.

Supported formats:
  - PDF  → pdfminer.six
  - DOCX → python-docx
"""
import os
import re
import logging

logger = logging.getLogger(__name__)


def parse_resume(file_path: str) -> str:
    """
    Parse a resume file (PDF or DOCX) and return extracted plain text.

    Args:
        file_path: Absolute path to the resume file.

    Returns:
        Cleaned plain text extracted from the resume.

    Raises:
        ValueError: If the file format is not supported.
        FileNotFoundError: If the file does not exist.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        text = _parse_pdf(file_path)
    elif ext == "docx":
        text = _parse_docx(file_path)
    elif ext == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Use PDF, DOCX, or TXT.")

    return _clean_text(text)


def _parse_pdf(file_path: str) -> str:
    """Extract text from a PDF using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(file_path)
        return text or ""
    except Exception as e:
        logger.error(f"PDF parsing failed for {file_path}: {e}")
        return ""


def _parse_docx(file_path: str) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        # Also read table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed for {file_path}: {e}")
        return ""


def _clean_text(text: str) -> str:
    """
    Clean raw extracted text:
    - Normalize whitespace and line breaks
    - Remove special characters (keep alphanumerics, basic punctuation)
    - Collapse multiple spaces
    """
    if not text:
        return ""

    # Replace non-breaking spaces and tabs
    text = text.replace("\xa0", " ").replace("\t", " ")
    # Normalize line endings
    text = re.sub(r"\r\n|\r", "\n", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)

    return text.strip()
